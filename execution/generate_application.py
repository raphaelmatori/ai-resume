"""
Application Generation Agent (Stage 3)
--------------------------------------
Parallel document generation script. Uses fact traceability to draft tailored
Resume and Cover Letter documents in .docx format.
"""
import os
import sys
# Add project root to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import logging
import concurrent.futures
from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import JsonOutputParser
from execution.utils import get_llm, ensure_directory, invoke_with_retry
from execution.gen_models import ResumeContent, CoverLetterContent

# Configure logging to stdout
logging.basicConfig(
    level=logging.INFO, 
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)

def load_file(path: str) -> str:
    if not os.path.exists(path):
        return ""
    with open(path, "r") as f:
        return f.read()

def strip_tags(text: str) -> str:
    """Removes <!-- source_id: ... --> tags from text."""
    if not text:
        return ""
    return re.sub(r"<!-- source_id:.*?-->", "", text).strip()

def generate_resume_content(candidate_md: str, vacancy_md: str) -> ResumeContent:
    llm = get_llm(temperature=0.1)
    parser = JsonOutputParser(pydantic_object=ResumeContent)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are a Professional Resume Strategist.

Context:
1. Candidate Profile (Strict Fact Database)
2. Vacancy Analysis (Target)

Task: Adapt the candidate's profile to the vacancy.

CRITICAL CONSTRAINTS:
- You may ONLY use facts present in the Candidate Profile.
- **Traceability**: Every bullet point or sentence MUST end with the original `<!-- source_id: ... -->` tag from the Candidate Profile.
- **STRICT FACTUAL CONSISTENCY**: 
  * If the profile says "8+ years", you MUST write "8+ years" - NOT "over 8 years", "6+ years", or any variation
  * If the profile says specific numbers (e.g., "~40k users"), preserve the EXACT phrasing
  * DO NOT round, approximate, or rephrase quantifiable facts
  * DO NOT calculate or derive new numbers from existing facts
  * Preserve exact dates, percentages, and numerical achievements
- If a section (like Projects or Certifications) has NO data in the profile, return an empty list for that section.
- **Presentation ONLY**: You may restructure how facts are presented, combine related bullets, or adjust emphasis - but NEVER change the facts themselves.

Output: serialized JSON matching the schema.
{format_instructions}
"""),
        ("human", "CANDIDATE PROFILE:\n{candidate}\n\nVACANCY:\n{vacancy}")
    ])
    
    chain = prompt | llm | parser
    logging.info("Generating tailored resume content...")
    return invoke_with_retry(chain, {
        "candidate": candidate_md,
        "vacancy": vacancy_md,
        "format_instructions": parser.get_format_instructions()
    }, max_retries=5)

def generate_cover_letter_content(candidate_md: str, vacancy_md: str) -> CoverLetterContent:
    llm = get_llm(temperature=0.2)
    parser = JsonOutputParser(pydantic_object=CoverLetterContent)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert career coach writing a modern, high-impact Cover Letter.

Task: Write a concise, punchy cover letter that connects the candidate's actual experience to the vacancy's needs.

Rules:
- **Brevity is Key**: Recruiters are busy. Keep the total length around 250-300 words.
- **Tone**: Professional, confident, and direct. Avoid generic "I am writing to..." fluff where possible. 
- **Evidence**: Focus on the 2-3 most relevant facts from the Candidate Profile that match the Vacancy.
- **Constraint**: Include `<!-- source_id: ... -->` comments in the text where facts are used.
- **STRICT FACTUAL CONSISTENCY**: 
  * Years of experience MUST match the Candidate Profile exactly (e.g., if "8+ years", write "8+ years" NOT "over 8 years" or "6+ years")
  * Preserve exact phrasing for quantifiable facts (numbers, percentages, user counts, etc.)
  * DO NOT invent, round, approximate, or modify any verifiable information
  * Example: "~40k users" stays "~40k users", NOT "40,000 users" or "approximately 40k users"
- **Greeting**: Always start with a professional greeting (e.g., "Dear Hiring Team at [Company Name],").
- **Signature**: Provide a professional sign-off (e.g. "Sincerely," or "Best regards,") followed by the candidate's full name in the `signature_name` field.
{format_instructions}
"""),
        ("human", "CANDIDATE PROFILE:\n{candidate}\n\nVACANCY:\n{vacancy}")
    ])
    
    chain = prompt | llm | parser
    logging.info("Generating tailored cover letter...")
    return invoke_with_retry(chain, {
        "candidate": candidate_md,
        "vacancy": vacancy_md,
        "format_instructions": parser.get_format_instructions()
    }, max_retries=5)

def create_docx(content: ResumeContent, output_path: str):
    doc = Document()
    
    # Global Style Adjustments
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    
    # Header
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = name_p.add_run(strip_tags(content.name).upper())
    run.bold = True
    run.font.size = Pt(16)

    role_p = doc.add_paragraph()
    role_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = role_p.add_run(strip_tags(content.role_title).upper())
    run.bold = True
    run.font.size = Pt(12)

    contact_p = doc.add_paragraph()
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_str = " | ".join([strip_tags(c) for c in content.contact_info if c])
    contact_p.add_run(contact_str)

    def add_section_header(title):
        header = doc.add_paragraph()
        header.paragraph_format.space_before = Pt(12)
        header.paragraph_format.space_after = Pt(6)
        run = header.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(11)

    # Summary
    if content.summary and strip_tags(content.summary):
        add_section_header("SUMMARY / PROFILE")
        doc.add_paragraph(strip_tags(content.summary))

    # Skills
    valid_skills = [strip_tags(s) for s in content.skills_section if strip_tags(s)]
    if valid_skills:
        add_section_header("TECHNICAL SKILLS")
        # For technical skills, comma separated looks better sometimes, but let's stick to a clean block
        # Grouped skills like in the example is harder with current schema, but we can try to join them
        doc.add_paragraph(", ".join(valid_skills))

    # Experience
    if content.experience_sections:
        add_section_header("PROFESSIONAL EXPERIENCE")
        for section in content.experience_sections:
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            run = title_p.add_run(strip_tags(section.title))
            run.bold = True
            for bullet in section.content:
                stripped = strip_tags(bullet)
                if stripped:
                    p = doc.add_paragraph(stripped, style='List Bullet')
                    p.paragraph_format.space_after = Pt(2)

    # Education
    valid_edu = [strip_tags(e) for e in content.education_section if strip_tags(e)]
    if valid_edu:
        add_section_header("EDUCATION")
        for edu in valid_edu:
            doc.add_paragraph(edu, style='List Bullet')

    # Certifications
    # Note: Check if certifications_section exists in content (Pydantic might not have it yet if not updated)
    certs = getattr(content, 'certifications_section', [])
    valid_certs = [strip_tags(c) for c in certs if strip_tags(c)]
    if valid_certs:
        add_section_header("CERTIFICATIONS")
        for cert in valid_certs:
            doc.add_paragraph(cert, style='List Bullet')

    # Projects
    valid_projs = [strip_tags(p) for p in content.projects_section if strip_tags(p)]
    if valid_projs:
        add_section_header("PROJECT HIGHLIGHTS")
        for proj in valid_projs:
            doc.add_paragraph(proj, style='List Bullet')

    doc.save(output_path)
    logging.info(f"Saved resume DOCX to {output_path}")

def create_cl_docx(content: CoverLetterContent, output_path: str):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Opening (greeting)
    opening_p = doc.add_paragraph(strip_tags(content.opening))
    opening_p.paragraph_format.space_after = Pt(12)  # Use spacing property instead of blank paragraph
    
    # Body paragraphs
    for para in content.body_paragraphs:
        stripped = strip_tags(para)
        if stripped:
            p = doc.add_paragraph(stripped)
            p.paragraph_format.space_after = Pt(12)  # Consistent spacing between paragraphs
    
    # Closing
    closing_p = doc.add_paragraph(strip_tags(content.closing))
    closing_p.paragraph_format.space_after = Pt(12)
    
    # Signature
    sig = strip_tags(content.signature_name)
    if sig:
        doc.add_paragraph(sig)
    
    doc.save(output_path)
    logging.info(f"Saved cover letter DOCX to {output_path}")

def main():
    # Use current working directory (important for packaged Electron app)
    BASE_DIR = os.getcwd()
    DATA_DIR = os.path.join(BASE_DIR, "data", "processed")
    OUTPUT_DIR = os.path.join(BASE_DIR, "output")
    ensure_directory(OUTPUT_DIR)
    
    candidate_path = os.path.join(DATA_DIR, "candidate_profile.md")
    vacancy_path = os.path.join(DATA_DIR, "vacancy_profile.md")
    
    if not os.path.exists(candidate_path) or not os.path.exists(vacancy_path):
        logging.error("Missing input files. Run ingestion scripts first.")
        return

    candidate_md = load_file(candidate_path)
    vacancy_md = load_file(vacancy_path)
    
    # Run Generation in Parallel to save time
    logging.info("Starting parallel generation of Resume and Cover Letter...")
    with concurrent.futures.ThreadPoolExecutor() as executor:
        resume_task = executor.submit(generate_resume_content, candidate_md, vacancy_md)
        cl_task = executor.submit(generate_cover_letter_content, candidate_md, vacancy_md)
        
        try:
            # 1. Process Resume
            resume_content = resume_task.result()
            if isinstance(resume_content, dict):
                resume_content = ResumeContent(**resume_content)
            create_docx(resume_content, os.path.join(OUTPUT_DIR, "Tailored_Resume.docx"))
            
            # 2. Process Cover Letter
            cl_content = cl_task.result()
            if isinstance(cl_content, dict):
                 cl_content = CoverLetterContent(**cl_content)
            create_cl_docx(cl_content, os.path.join(OUTPUT_DIR, "Tailored_CoverLetter.docx"))
            
        except Exception as e:
            logging.error(f"Generation failed: {e}")
            sys.exit(1)

if __name__ == "__main__":
    main()
